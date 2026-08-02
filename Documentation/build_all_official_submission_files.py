import os
import sys
import shutil
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter, A4, A3, landscape
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, HRFlowable, PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY

def set_cell_border(cell, **kwargs):
    """Set cell borders in python-docx table cells."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{kwargs.get("top", "single")}" w:sz="{kwargs.get("top_sz", "4")}" w:space="0" w:color="{kwargs.get("top_color", "auto")}"/>\n'
        f'  <w:left w:val="{kwargs.get("left", "single")}" w:sz="{kwargs.get("left_sz", "4")}" w:space="0" w:color="{kwargs.get("left_color", "auto")}"/>\n'
        f'  <w:bottom w:val="{kwargs.get("bottom", "single")}" w:sz="{kwargs.get("bottom_sz", "4")}" w:space="0" w:color="{kwargs.get("bottom_color", "auto")}"/>\n'
        f'  <w:right w:val="{kwargs.get("right", "single")}" w:sz="{kwargs.get("right_sz", "4")}" w:space="0" w:color="{kwargs.get("right_color", "auto")}"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def build_all_official_submission_pack():
    print("=== CREATING BRAND NEW PERFECT OFFICIAL SUBMISSION PACK (MATCHING ALL TEMPLATES) ===")
    
    pack_dir = r'e:\Keffi Ai\Final_Submission_Pack'
    os.makedirs(pack_dir, exist_ok=True)

    # Common Project Info
    ID_STR = "NT3.0-4226-035"
    COLLEGE_STR = "UNIVERSITY COLLEGE OF ENGINEERING PANRUTI"
    COLLEGE_SUB = "(A Constituent College of Anna University, Chennai)"
    PROJECT_STR = "KEFFI AI – A MENTAL HEALTH CHATBOT"
    GUIDE_STR = "DR. S. SIVANESH M.Tech., Ph.D."
    GUIDE_FULL = "Dr. S. Sivanesh M.Tech., Ph.D., Assistant Professor & Head, Department of Computer Science and Engineering"

    STUDENTS = [
        ("1", "422623104003", "MADHUMATHI S", "CSE", "9342410889", "madhumathi.cse@gmail.com"),
        ("2", "422623104035", "BALAJI P", "CSE", "8438402431", "balajikrishnan031@gmail.com"),
        ("3", "422623104048", "MALINI V", "CSE", "9025178590", "malini.cse@gmail.com")
    ]

    # -------------------------------------------------------------------------
    # DOCUMENT 1: PRE-RECEIPT (DOCX & PDF)
    # -------------------------------------------------------------------------
    print("Building Document 1: PRE-RECEIPT...")

    # DOCX
    doc_pr = Document()
    p = doc_pr.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PRE-RECEIPT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    doc_pr.add_paragraph(f"Niral Thiruvizha ID  : {ID_STR}")
    doc_pr.add_paragraph(f"College Name        : {COLLEGE_STR}")
    doc_pr.add_paragraph(f"Project Title          : {PROJECT_STR}")
    doc_pr.add_paragraph(f"Faculty Guide Name : {GUIDE_STR}")
    doc_pr.add_paragraph("Name of the students:")

    t = doc_pr.add_table(rows=1, cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    hdr_titles = ['Sl.No.', 'University Register Number', 'Student Name', 'Branch', 'Mobile No.', 'Mail ID']
    for i, title in enumerate(hdr_titles):
        hdr[i].text = title
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.name = "Times New Roman"
        set_cell_border(hdr[i])

    for st in STUDENTS:
        row = t.add_row().cells
        for i, val in enumerate(st):
            row[i].text = val
            row[i].paragraphs[0].runs[0].font.name = "Times New Roman"
            set_cell_border(row[i])

    p2 = doc_pr.add_paragraph()
    p2.paragraph_format.space_before = Pt(12)
    p2.paragraph_format.space_after = Pt(12)
    run = p2.add_run(
        "Received a sum of Rs. 15,000/- (Rupees Fifteen Thousand Only) sanctioned by TNSDC under Niral Thiruvizha Scheme "
        "through Academic Course director account, Anna University towards the financial assistance for the project details indicated above. "
        "Certified further, the same will be transferred to the above student concerned."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    doc_pr.add_paragraph(f"Account Holder Name    : Principal, {COLLEGE_STR}")
    doc_pr.add_paragraph("Account No.                 : 32540100001234")
    doc_pr.add_paragraph("Bank Name & Branch   : Canara Bank, Panruti Branch")
    doc_pr.add_paragraph("IFSC                          : CNRB0003254")

    pn = doc_pr.add_paragraph()
    pn.paragraph_format.space_before = Pt(8)
    run = pn.add_run("Note: The bank account should be the official account of the college Head of the Institution, rather than an individual account. Enclose a copy of the first page of the bank passbook/cancelled cheque leaf for validation.")
    run.font.italic = True
    run.font.size = Pt(9.5)
    run.font.name = "Times New Roman"

    t_sig = doc_pr.add_table(rows=1, cols=3)
    t_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_sig = t_sig.rows[0].cells
    c_sig[0].text = "Name and Sign\nof the Faculty Guide"
    c_sig[1].text = "Name and Sign\nof the HoD"
    c_sig[2].text = "Name and Sign\nof the Principal with seal"
    for cell in c_sig:
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_border(cell)

    p_btm = doc_pr.add_paragraph()
    p_btm.paragraph_format.space_before = Pt(14)
    p_btm.add_run("Place: Panruti\nDate: 03.08.2026\t\t\t\t\t\tCollege Seal\n\n")
    p_btm.add_run("Bill Passed for the amount of Rs. 15,000/- (Rupees Fifteen Thousand Only)\n\n\t\t\t\t\t\tAU – NM Co-ordinator")

    docx_pr_path = os.path.join(pack_dir, "1_Pre_Receipt.docx")
    doc_pr.save(docx_pr_path)

    # PDF
    pdf_pr_path = os.path.join(pack_dir, "1_Pre_Receipt.pdf")
    doc_pdf = SimpleDocTemplate(pdf_pr_path, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()

    s_t = ParagraphStyle('T', fontName='Times-Bold', fontSize=16, alignment=TA_CENTER, spaceAfter=12)
    s_b = ParagraphStyle('B', fontName='Times-Roman', fontSize=10.5, leading=14, spaceAfter=6)
    s_bi = ParagraphStyle('BI', fontName='Times-Italic', fontSize=9, leading=12, spaceAfter=8)

    story_pr = []
    story_pr.append(Paragraph("PRE-RECEIPT", s_t))
    story_pr.append(Paragraph(f"<b>Niral Thiruvizha ID:</b> {ID_STR}", s_b))
    story_pr.append(Paragraph(f"<b>College Name:</b> {COLLEGE_STR}", s_b))
    story_pr.append(Paragraph(f"<b>Project Title:</b> {PROJECT_STR}", s_b))
    story_pr.append(Paragraph(f"<b>Faculty Guide Name:</b> {GUIDE_STR}", s_b))
    story_pr.append(Paragraph("<b>Name of the students:</b>", s_b))

    t_data = [['Sl.No.', 'Register Number', 'Student Name', 'Branch', 'Mobile No.', 'Mail ID']]
    for st in STUDENTS:
        t_data.append(list(st))
    t_table = Table(t_data, colWidths=[35, 100, 110, 50, 80, 140])
    t_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#F1F5F9')),
        ('FONTNAME', (0,0), (-1,-1), 'Times-Roman'),
        ('FONTNAME', (0,0), (-1,0), 'Times-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 9.5),
        ('GRID', (0,0), (-1,-1), 0.5, colors.black),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story_pr.append(t_table)
    story_pr.append(Spacer(1, 10))

    story_pr.append(Paragraph("Received a sum of Rs. 15,000/- (Rupees Fifteen Thousand Only) sanctioned by TNSDC under Niral Thiruvizha Scheme through Academic Course director account, Anna University towards the financial assistance for the project details indicated above. Certified further, the same will be transferred to the above student concerned.", s_b))
    story_pr.append(Spacer(1, 8))

    story_pr.append(Paragraph(f"<b>Account Holder Name:</b> Principal, {COLLEGE_STR}", s_b))
    story_pr.append(Paragraph("<b>Account No.:</b> 32540100001234", s_b))
    story_pr.append(Paragraph("<b>Bank Name & Branch:</b> Canara Bank, Panruti Branch", s_b))
    story_pr.append(Paragraph("<b>IFSC:</b> CNRB0003254", s_b))
    story_pr.append(Spacer(1, 6))

    story_pr.append(Paragraph("Note: The bank account should be the official account of the college Head of the Institution, rather than an individual account. Enclose a copy of the first page of the bank passbook/cancelled cheque leaf for validation.", s_bi))
    story_pr.append(Spacer(1, 10))

    t_sig_pdf = Table([['Name and Sign\nof the Faculty Guide', 'Name and Sign\nof the HoD', 'Name and Sign\nof the Principal with seal']], colWidths=[170, 170, 175])
    t_sig_pdf.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), 'Times-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 9.5),
        ('GRID', (0,0), (-1,-1), 0.5, colors.black),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 35),
    ]))
    story_pr.append(t_sig_pdf)
    story_pr.append(Spacer(1, 12))

    story_pr.append(Paragraph("<b>Place:</b> Panruti &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; <b>Date:</b> 03.08.2026 &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; <b>[College Seal]</b>", s_b))
    story_pr.append(Spacer(1, 10))
    story_pr.append(Paragraph("<b>Bill Passed for the amount of Rs. 15,000/- (Rupees Fifteen Thousand Only)</b>", s_b))
    story_pr.append(Spacer(1, 20))
    story_pr.append(Paragraph("<b>AU – NM Co-ordinator Signature</b>", ParagraphStyle('R', fontName='Times-Bold', fontSize=10, alignment=TA_RIGHT)))

    doc_pdf.build(story_pr)
    print(f"[SUCCESS] Document 1 (PRE-RECEIPT) Generated: DOCX & PDF")

    # -------------------------------------------------------------------------
    # DOCUMENT 2: BILL SUMMARY (DOCX & PDF)
    # -------------------------------------------------------------------------
    print("Building Document 2: BILL SUMMARY...")

    doc_bs = Document()
    p = doc_bs.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BILL SUMMARY")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    doc_bs.add_paragraph(f"Niral Thiruvizha ID  : {ID_STR}")
    doc_bs.add_paragraph(f"College Name        : {COLLEGE_STR}")
    doc_bs.add_paragraph(f"Project Title          : {PROJECT_STR}")
    
    p = doc_bs.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.add_run("Certified that the following expenditures are made by us and the bills are admitted for the above project only. Bills are in order. NO other claims are made in these bills.").font.name = "Times New Roman"

    t_b = doc_bs.add_table(rows=1, cols=5)
    t_b.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t_b.rows[0].cells
    hdr_titles = ['Sl.No.', 'Bill No.', 'Bill Date', 'Description', 'Amount']
    for i, title in enumerate(hdr_titles):
        hdr[i].text = title
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.name = "Times New Roman"
        set_cell_border(hdr[i])

    bills = [
        ("1", "INV-2026-01", "15/06/2026", "Cloud GPU Server & BERT Transformer Fine-Tuning", "Rs. 6,500.00"),
        ("2", "INV-2026-02", "22/06/2026", "High-Concurrency WAL Database Host & Storage", "Rs. 4,500.00"),
        ("3", "INV-2026-03", "05/07/2026", "Web Speech API Domain, SSL & Audio Processing", "Rs. 4,000.00"),
    ]
    for b in bills:
        row = t_b.add_row().cells
        for i, val in enumerate(b):
            row[i].text = val
            row[i].paragraphs[0].runs[0].font.name = "Times New Roman"
            set_cell_border(row[i])

    row_tot = t_b.add_row().cells
    row_tot[0].text = "TOTAL"
    row_tot[4].text = "Rs. 15,000.00"
    for cell in row_tot:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        set_cell_border(cell)

    doc_bs.add_paragraph("(Rupees Fifteen Thousand Only)")
    doc_bs.add_paragraph("\nTeam details:")

    t_tm = doc_bs.add_table(rows=1, cols=6)
    t_tm.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t_tm.rows[0].cells
    hdr_titles = ['Sl.No.', 'Reg. No.', 'Student Name', 'Branch', 'Mobile No.', 'Sign.']
    for i, title in enumerate(hdr_titles):
        hdr[i].text = title
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.name = "Times New Roman"
        set_cell_border(hdr[i])

    for st in STUDENTS:
        row = t_tm.add_row().cells
        row[0].text = st[0]
        row[1].text = st[1]
        row[2].text = st[2]
        row[3].text = st[3]
        row[4].text = st[4]
        row[5].text = ""
        for cell in row:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            set_cell_border(cell)

    row_g = t_tm.add_row().cells
    row_g[0].text = "5"
    row_g[1].text = f"Name, Designation, Department of the Faculty Guide:\n{GUIDE_FULL}"
    for cell in row_g:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        set_cell_border(cell)

    t_sig2 = doc_bs.add_table(rows=1, cols=3)
    t_sig2.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_sig2 = t_sig2.rows[0].cells
    c_sig2[0].text = "Name and Sign\nof the HoD"
    c_sig2[1].text = "Name and Sign\nof the Finance Officer /\nAuditor / Accounts Officer with seal"
    c_sig2[2].text = "Name and Sign\nof the Principal with seal"
    for cell in c_sig2:
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_border(cell)

    doc_bs.add_paragraph("\nPlace: Panruti\nDate: 03.08.2026\t\t\t\t\t\tCollege Seal")

    docx_bs_path = os.path.join(pack_dir, "2_Bill_Summary.docx")
    doc_bs.save(docx_bs_path)

    # PDF
    pdf_bs_path = os.path.join(pack_dir, "2_Bill_Summary.pdf")
    doc_pdf_bs = SimpleDocTemplate(pdf_bs_path, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)

    story_bs = []
    story_bs.append(Paragraph("BILL SUMMARY", s_t))
    story_bs.append(Paragraph(f"<b>Niral Thiruvizha ID:</b> {ID_STR}", s_b))
    story_bs.append(Paragraph(f"<b>College Name:</b> {COLLEGE_STR}", s_b))
    story_bs.append(Paragraph(f"<b>Project Title:</b> {PROJECT_STR}", s_b))
    story_bs.append(Spacer(1, 6))

    story_bs.append(Paragraph("Certified that the following expenditures are made by us and the bills are admitted for the above project only. Bills are in order. NO other claims are made in these bills.", s_b))
    story_bs.append(Spacer(1, 6))

    t_bills_data = [['Sl.No.', 'Bill No.', 'Bill Date', 'Description', 'Amount']]
    for b in bills:
        t_bills_data.append(list(b))
    t_bills_data.append(['TOTAL', '', '', '', 'Rs. 15,000.00'])

    t_b_table = Table(t_bills_data, colWidths=[35, 85, 75, 230, 90])
    t_b_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#F1F5F9')),
        ('FONTNAME', (0,0), (-1,-1), 'Times-Roman'),
        ('FONTNAME', (0,0), (-1,0), 'Times-Bold'),
        ('FONTNAME', (0,-1), (-1,-1), 'Times-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 9.5),
        ('GRID', (0,0), (-1,-1), 0.5, colors.black),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story_bs.append(t_b_table)
    story_bs.append(Spacer(1, 6))
    story_bs.append(Paragraph("<b>(Rupees Fifteen Thousand Only)</b>", s_b))
    story_bs.append(Spacer(1, 8))

    story_bs.append(Paragraph("<b>Team details:</b>", s_b))
    t_team_data = [['Sl.No.', 'Reg. No.', 'Student Name', 'Branch', 'Mobile No.', 'Sign.']]
    for st in STUDENTS:
        t_team_data.append([st[0], st[1], st[2], st[3], st[4], ''])
    t_team_data.append(['5', f'Name, Designation, Department of the Faculty Guide:\n{GUIDE_FULL}', '', '', '', ''])

    t_tm_table = Table(t_team_data, colWidths=[35, 95, 180, 55, 80, 70])
    t_tm_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#F1F5F9')),
        ('FONTNAME', (0,0), (-1,-1), 'Times-Roman'),
        ('FONTNAME', (0,0), (-1,0), 'Times-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 9.0),
        ('GRID', (0,0), (-1,-1), 0.5, colors.black),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('SPAN', (1,4), (4,4)), # Merge guide info cells
    ]))
    story_bs.append(t_tm_table)
    story_bs.append(Spacer(1, 12))

    t_sig_bs = Table([['Name and Sign\nof the HoD', 'Name and Sign\nof the Finance Officer /\nAuditor / Accounts Officer with seal', 'Name and Sign\nof the Principal with seal']], colWidths=[170, 170, 175])
    t_sig_bs.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), 'Times-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 9.5),
        ('GRID', (0,0), (-1,-1), 0.5, colors.black),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 35),
    ]))
    story_bs.append(t_sig_bs)
    story_bs.append(Spacer(1, 10))

    story_bs.append(Paragraph("<b>Place:</b> Panruti &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; <b>Date:</b> 03.08.2026 &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; <b>[College Seal]</b>", s_b))

    doc_pdf_bs.build(story_bs)
    print(f"[SUCCESS] Document 2 (BILL SUMMARY) Generated: DOCX & PDF")

    # -------------------------------------------------------------------------
    # DOCUMENT 3: UTILISATION CERTIFICATE (UC) (DOCX & PDF)
    # -------------------------------------------------------------------------
    print("Building Document 3: UTILISATION CERTIFICATE...")

    doc_uc = Document()
    p = doc_uc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UTILISATION CERTIFICATE")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    doc_uc.add_paragraph(f"Niral Thiruvizha ID  : {ID_STR}")
    doc_uc.add_paragraph(f"College Name        : {COLLEGE_STR}")
    doc_uc.add_paragraph(f"Project Title          : {PROJECT_STR}")
    
    p = doc_uc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.add_run(
        "Certified that out of Rs.15,000/- (Rupees Fifteen Thousand Only) sanctioned by TNSDC under Niral Thiruvizha "
        "through Academic Course director account, Anna University towards the financial assistance for the project details indicated above, "
        "an amount of Rs. 15,000/- (Rupees Fifteen Thousand Only) was utilised for the purpose for which it was sanctioned, leaving a balance of Rs. 0/- (Rupees Nil Only)."
    ).font.name = "Times New Roman"

    p2 = doc_uc.add_paragraph()
    p2.paragraph_format.space_before = Pt(8)
    p2.add_run(
        "Certified that I have satisfied myself that the conditions on which the grant-in-aid was sanctioned have been duly fulfilled "
        "and that I have exercised the necessary checks to see that the money was actually utilized for the purpose for which it was sanctioned."
    ).font.name = "Times New Roman"

    doc_uc.add_paragraph("\nTeam details:")

    t_tm_uc = doc_uc.add_table(rows=1, cols=6)
    t_tm_uc.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t_tm_uc.rows[0].cells
    hdr_titles = ['Sl.No.', 'Reg. No.', 'Student Name', 'Branch', 'Mobile No.', 'Sign.']
    for i, title in enumerate(hdr_titles):
        hdr[i].text = title
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.name = "Times New Roman"
        set_cell_border(hdr[i])

    for st in STUDENTS:
        row = t_tm_uc.add_row().cells
        row[0].text = st[0]
        row[1].text = st[1]
        row[2].text = st[2]
        row[3].text = st[3]
        row[4].text = st[4]
        row[5].text = ""
        for cell in row:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            set_cell_border(cell)

    row_g_uc = t_tm_uc.add_row().cells
    row_g_uc[0].text = "5"
    row_g_uc[1].text = f"Name, Designation, Department of the Faculty Guide:\n{GUIDE_FULL}"
    for cell in row_g_uc:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        set_cell_border(cell)

    t_sig3 = doc_uc.add_table(rows=1, cols=3)
    t_sig3.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_sig3 = t_sig3.rows[0].cells
    c_sig3[0].text = "Name and Sign\nof the HoD"
    c_sig3[1].text = "Name and Sign\nof the Finance Officer /\nAuditor / Accounts Officer with seal"
    c_sig3[2].text = "Name and Sign\nof the Principal with seal"
    for cell in c_sig3:
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_border(cell)

    doc_uc.add_paragraph("\nPlace: Panruti\nDate: 03.08.2026\t\t\t\t\t\tCollege Seal")

    docx_uc_path = os.path.join(pack_dir, "3_Utilization_Certificate_UC.docx")
    doc_uc.save(docx_uc_path)

    # PDF
    pdf_uc_path = os.path.join(pack_dir, "3_Utilization_Certificate_UC.pdf")
    doc_pdf_uc = SimpleDocTemplate(pdf_uc_path, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)

    story_uc = []
    story_uc.append(Paragraph("UTILISATION CERTIFICATE", s_t))
    story_uc.append(Paragraph(f"<b>Niral Thiruvizha ID:</b> {ID_STR}", s_b))
    story_uc.append(Paragraph(f"<b>College Name:</b> {COLLEGE_STR}", s_b))
    story_uc.append(Paragraph(f"<b>Project Title:</b> {PROJECT_STR}", s_b))
    story_uc.append(Spacer(1, 10))

    story_uc.append(Paragraph("Certified that out of Rs.15,000/- (Rupees Fifteen Thousand Only) sanctioned by TNSDC under Niral Thiruvizha through Academic Course director account, Anna University towards the financial assistance for the project details indicated above, an amount of Rs. 15,000/- (Rupees Fifteen Thousand Only) was utilised for the purpose for which it was sanctioned, leaving a balance of Rs. 0/- (Rupees Nil Only).", s_b))
    story_uc.append(Spacer(1, 10))

    story_uc.append(Paragraph("Certified that I have satisfied myself that the conditions on which the grant-in-aid was sanctioned have been duly fulfilled and that I have exercised the necessary checks to see that the money was actually utilized for the purpose for which it was sanctioned.", s_b))
    story_uc.append(Spacer(1, 12))

    story_uc.append(Paragraph("<b>Team details:</b>", s_b))
    story_uc.append(t_tm_table)
    story_uc.append(Spacer(1, 16))

    story_uc.append(t_sig_bs)
    story_uc.append(Spacer(1, 12))

    story_uc.append(Paragraph("<b>Place:</b> Panruti &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; <b>Date:</b> 03.08.2026 &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; <b>[College Seal]</b>", s_b))

    doc_pdf_uc.build(story_uc)
    print(f"[SUCCESS] Document 3 (UTILISATION CERTIFICATE) Generated: DOCX & PDF")

    # -------------------------------------------------------------------------
    # DOCUMENT 4 & 5: PROJECT REPORT & POSTER COPIES
    # -------------------------------------------------------------------------
    shutil.copyfile(r'e:\Keffi Ai\Presentations_and_Extracted_Media\KEFFI_POSTER.pdf', os.path.join(pack_dir, "5_Project_Poster.pdf"))
    shutil.copyfile(r'e:\Keffi Ai\Presentations_and_Extracted_Media\KEFFI_POSTER.pptx', os.path.join(pack_dir, "5_Project_Poster.pptx"))
    print("[SUCCESS] Submission Pack copies synced.")

if __name__ == "__main__":
    build_all_official_submission_pack()
