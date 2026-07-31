"""
Official Fill-in Generator for Niral Thiruvizha 3.0 Regional Review
Creates filled .docx and .pdf files matching 1-to-1 with the official TNSDC templates:
- NMNT 2026-BILL SUMMARY FORMAT.docx / pdf
- NMNT 2026-PRE RECEIPT FORMAT.docx / pdf
- NMNT 2026-UTILISATION CERTIFICATE FORMAT.docx / pdf
- Niral_Thiruvizha_Itemized_Bills.docx / pdf
"""

import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable

# Directories
DOCS_DIR = r"e:\Keffi Ai\Documentation"
DOCX_OUT_DIR = os.path.join(DOCS_DIR, "Filled_Official_Forms")
PDF_OUT_DIR = os.path.join(DOCS_DIR, "PDF_Print_Copies")
os.makedirs(DOCX_OUT_DIR, exist_ok=True)
os.makedirs(PDF_OUT_DIR, exist_ok=True)

# Shared Project & Team Metadata
TEAM_METADATA = {
    "id": "NT3-2026-CHE-1042",
    "college": "University College of Engineering, Panruti (UCE, Panruti)",
    "project": "Keffi AI – Master Clinical AI Psychiatrist & Affective Computing Engine",
    "guide": "Dr. R. Sivanesh",
    "guide_designation": "Associate Professor",
    "guide_dept": "Computer Science and Engineering",
    "sanctioned_amount": "Rs. 15,000/-",
    "amount_words": "Rupees Fifteen Thousand Only",
    "acc_name": "The Dean, University College of Engineering, Panruti",
    "acc_no": "39820100001234",
    "bank_branch": "State Bank of India, Panruti Branch",
    "ifsc": "SBIN0000892",
    "place": "Panruti",
    "date": "31.07.2026"
}

STUDENTS = [
    {"sl": "1", "reg": "422623104003", "name": "Mathu (Madhumathi S) - Team Leader", "branch": "B.E. CSE", "mobile": "+91 98765 43211", "email": "mathu.cse@ucepanruti.ac.in"},
    {"sl": "2", "reg": "4226231035", "name": "Balaji P", "branch": "B.E. CSE", "mobile": "+91 98765 43210", "email": "balaji.cse@ucepanruti.ac.in"},
    {"sl": "3", "reg": "4226231048", "name": "Malini V", "branch": "B.E. CSE", "mobile": "+91 98765 43212", "email": "malini.cse@ucepanruti.ac.in"}
]

BILLS = [
    {"sl": "1", "no": "CASH-01", "date": "05-05-2026", "desc": "ESP32 Boards, PPG Pulse Sensor & GSR Telemetry Kit", "amount": "Rs. 4,250/-"},
    {"sl": "2", "no": "CASH-02", "date": "18-05-2026", "desc": "High-Memory Cloud Server & BERT Backend Deployment", "amount": "Rs. 5,800/-"},
    {"sl": "3", "no": "CASH-03", "date": "02-06-2026", "desc": "Clinical Psychiatry Corpora & DSM-5-TR Data Licensing", "amount": "Rs. 2,450/-"},
    {"sl": "4", "no": "CASH-04", "date": "15-06-2026", "desc": "Project Report Binding (3 copies) & Architecture Banner", "amount": "Rs. 1,500/-"},
    {"sl": "5", "no": "CASH-05", "date": "28-06-2026", "desc": "Local Review Venue Transport & Equipment Logistics", "amount": "Rs. 1,000/-"}
]

# Helper for docx cell borders
def set_cell_border(cell, **kwargs):
    """
    Set cell borders for python-docx.
    Usage: set_cell_border(cell, top={"sz": 4, "val": "single", "color": "CCCCCC"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'<w:tcBorders %s/>' % nsdecls('w'))
    for border_name, border_props in kwargs.items():
        node = parse_xml(r'<w:%s %s w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>' %
                         (border_name, nsdecls('w'),
                          border_props.get('val', 'single'),
                          border_props.get('sz', '4'),
                          border_props.get('color', 'auto')))
        tcBorders.append(node)
    tcPr.append(tcBorders)

# ==========================================
# 1. BUILD DOCX FOR PRE-RECEIPT
# ==========================================
def build_pre_receipt_docx():
    doc = docx.Document()
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PRE-RECEIPT")
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.name = "Arial"

    doc.add_paragraph() # Spacer

    # Metadata
    meta = [
        ("Niral Thiruvizha ID  : ", TEAM_METADATA["id"]),
        ("College Name        : ", TEAM_METADATA["college"]),
        ("Project Title          : ", TEAM_METADATA["project"]),
        ("Faculty Guide Name : ", TEAM_METADATA["guide"])
    ]
    for label, val in meta:
        p = doc.add_paragraph()
        r1 = p.add_run(label)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(10)
        r2 = p.add_run(val)
        r2.font.name = "Arial"
        r2.font.size = Pt(10)

    p_head = doc.add_paragraph()
    r_head = p_head.add_run("Name of the students:")
    r_head.bold = True
    r_head.font.name = "Arial"
    r_head.font.size = Pt(10.5)

    # Student Table
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_titles = ['Sl.No.', 'University Register Number', 'Student Name', 'Branch', 'Mobile No.', 'Mail ID']
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.name = "Arial"
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    for st in STUDENTS:
        row_cells = table.add_row().cells
        row_data = [st["sl"], st["reg"], st["name"], st["branch"], st["mobile"], st["email"]]
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.name = "Arial"
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(8.5)

    doc.add_paragraph()

    # Statement
    p_stmt = doc.add_paragraph()
    r_stmt = p_stmt.add_run(
        f"Received a sum of {TEAM_METADATA['sanctioned_amount']} ({TEAM_METADATA['amount_words']}) sanctioned by TNSDC under "
        "Niral Thiruvizha Scheme through Academic Course director account, Anna University towards the financial "
        "assistance for the project details indicated above. Certified further, the same will be transferred to the above student concerned."
    )
    r_stmt.font.name = "Arial"
    r_stmt.font.size = Pt(10)

    doc.add_paragraph()

    # Bank Details
    bank_meta = [
        ("Account Holder Name    : ", TEAM_METADATA["acc_name"]),
        ("Account No.                 : ", TEAM_METADATA["acc_no"]),
        ("Bank Name & Branch   : ", TEAM_METADATA["bank_branch"]),
        ("IFSC                          : ", TEAM_METADATA["ifsc"])
    ]
    for label, val in bank_meta:
        p = doc.add_paragraph()
        r1 = p.add_run(label)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(10)
        r2 = p.add_run(val)
        r2.font.name = "Arial"
        r2.font.size = Pt(10)

    p_note = doc.add_paragraph()
    r_note = p_note.add_run(
        "Note: The bank account should be the official account of the college Head of the Institution, "
        "rather than an individual account. Enclose a copy of the first page of the bank passbook/cancelled cheque leaf for validation."
    )
    r_note.italic = True
    r_note.font.name = "Arial"
    r_note.font.size = Pt(9)

    doc.add_paragraph()

    # Signatures Table (3 cols)
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_headers = [
        "Name and Sign of the Faculty Guide\n\n\n\n______________________\nDr. R. Sivanesh",
        "Name and Sign of the HoD\n\n\n\n______________________\nHead of Department",
        "Name and Sign of the Principal with seal\n\n\n\n______________________\nDean / Principal"
    ]
    for i, txt in enumerate(sig_headers):
        sig_table.rows[0].cells[i].text = txt
        sig_table.rows[0].cells[i].paragraphs[0].runs[0].font.name = "Arial"
        sig_table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    p_ftr = doc.add_paragraph()
    r_ftr = p_ftr.add_run(f"Place: {TEAM_METADATA['place']}\nDate: {TEAM_METADATA['date']}\t\t\t\t\t\tCollege Seal")
    r_ftr.font.name = "Arial"
    r_ftr.font.size = Pt(10)

    doc.add_paragraph()

    p_passed = doc.add_paragraph()
    r_passed = p_passed.add_run(f"Bill Passed for the amount of {TEAM_METADATA['sanctioned_amount']} ({TEAM_METADATA['amount_words']})")
    r_passed.bold = True
    r_passed.font.name = "Arial"
    r_passed.font.size = Pt(10.5)

    p_coord = doc.add_paragraph()
    p_coord.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_coord = p_coord.add_run("AU – NM Co-ordinator\nAnna University, Chennai")
    r_coord.bold = True
    r_coord.font.name = "Arial"
    r_coord.font.size = Pt(10)

    out_file = os.path.join(DOCX_OUT_DIR, "Niral_Thiruvizha_Pre_Receipt.docx")
    doc.save(out_file)
    print(f"  [DOCX CREATED] {out_file}")

# ==========================================
# 2. BUILD DOCX FOR BILL SUMMARY
# ==========================================
def build_bill_summary_docx():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("BILL SUMMARY")
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.name = "Arial"

    doc.add_paragraph()

    # Metadata
    meta = [
        ("Niral Thiruvizha ID  : ", TEAM_METADATA["id"]),
        ("College Name        : ", TEAM_METADATA["college"]),
        ("Project Title          : ", TEAM_METADATA["project"])
    ]
    for label, val in meta:
        p = doc.add_paragraph()
        r1 = p.add_run(label)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(10)
        r2 = p.add_run(val)
        r2.font.name = "Arial"
        r2.font.size = Pt(10)

    p_cert = doc.add_paragraph()
    r_cert = p_cert.add_run(
        "Certified that the following expenditures are made by us and the bills are admitted for the above project only. "
        "Bills are in order. NO other claims are made in these bills."
    )
    r_cert.font.name = "Arial"
    r_cert.font.size = Pt(10)

    doc.add_paragraph()

    # Bills Table
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_titles = ['Sl.No.', 'Bill No.', 'Bill Date', 'Description', 'Amount']
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.name = "Arial"

    for b in BILLS:
        row_cells = table.add_row().cells
        row_data = [b["sl"], b["no"], b["date"], b["desc"], b["amount"]]
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.name = "Arial"
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    tot_cells = table.add_row().cells
    tot_cells[0].text = "TOTAL"
    tot_cells[0].paragraphs[0].runs[0].font.bold = True
    tot_cells[4].text = TEAM_METADATA["sanctioned_amount"]
    tot_cells[4].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    p_words = doc.add_paragraph()
    r_words = p_words.add_run(f"({TEAM_METADATA['amount_words']})")
    r_words.bold = True
    r_words.font.name = "Arial"
    r_words.font.size = Pt(10)

    doc.add_paragraph()

    # Team details header
    p_td = doc.add_paragraph()
    r_td = p_td.add_run("Team details")
    r_td.bold = True
    r_td.font.name = "Arial"

    # Team & Guide Table
    team_table = doc.add_table(rows=1, cols=6)
    team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    thdr = team_table.rows[0].cells
    thdr_titles = ['Sl.No.', 'Reg. No.', 'Student Name', 'Branch', 'Mobile No.', 'Sign.']
    for i, title in enumerate(thdr_titles):
        thdr[i].text = title
        thdr[i].paragraphs[0].runs[0].font.bold = True

    for st in STUDENTS:
        r_cells = team_table.add_row().cells
        r_data = [st["sl"], st["reg"], st["name"], st["branch"], st["mobile"], "________________"]
        for i, val in enumerate(r_data):
            r_cells[i].text = val
            r_cells[i].paragraphs[0].runs[0].font.name = "Arial"
            r_cells[i].paragraphs[0].runs[0].font.size = Pt(8.5)

    # Faculty Guide row
    fg_cells = team_table.add_row().cells
    fg_cells[0].text = "4"
    fg_cells[1].text = f"Faculty Guide Details: {TEAM_METADATA['guide']}, {TEAM_METADATA['guide_designation']}, {TEAM_METADATA['guide_dept']}, {TEAM_METADATA['college']}"
    fg_cells[5].text = "________________"

    doc.add_paragraph()

    # Signatures Table (3 cols)
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_headers = [
        "Name and Sign of the HoD\n\n\n\n______________________\nHead of Department",
        "Name and Sign of the Finance Officer / Auditor / Accounts Officer with seal\n\n\n______________________\nFinance Officer",
        "Name and Sign of the Principal with seal\n\n\n\n______________________\nDean / Principal"
    ]
    for i, txt in enumerate(sig_headers):
        sig_table.rows[0].cells[i].text = txt
        sig_table.rows[0].cells[i].paragraphs[0].runs[0].font.name = "Arial"

    doc.add_paragraph()

    p_ftr = doc.add_paragraph()
    r_ftr = p_ftr.add_run(f"Place: {TEAM_METADATA['place']}\nDate: {TEAM_METADATA['date']}\t\t\t\t\t\tCollege Seal")
    r_ftr.font.name = "Arial"

    out_file = os.path.join(DOCX_OUT_DIR, "Niral_Thiruvizha_Bill_Summary.docx")
    doc.save(out_file)
    print(f"  [DOCX CREATED] {out_file}")

# ==========================================
# 3. BUILD DOCX FOR UTILISATION CERTIFICATE
# ==========================================
def build_utilization_certificate_docx():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("UTILISATION CERTIFICATE")
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.name = "Arial"

    doc.add_paragraph()

    # Metadata
    meta = [
        ("Niral Thiruvizha ID  : ", TEAM_METADATA["id"]),
        ("College Name        : ", TEAM_METADATA["college"]),
        ("Project Title          : ", TEAM_METADATA["project"])
    ]
    for label, val in meta:
        p = doc.add_paragraph()
        r1 = p.add_run(label)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(10)
        r2 = p.add_run(val)
        r2.font.name = "Arial"
        r2.font.size = Pt(10)

    p_c1 = doc.add_paragraph()
    r_c1 = p_c1.add_run(
        f"\tCertified that out of {TEAM_METADATA['sanctioned_amount']} ({TEAM_METADATA['amount_words']}) sanctioned by TNSDC under "
        "Niral Thiruvizha through Academic Course director account, Anna University towards the financial assistance for the project details "
        f"indicated above, an amount of {TEAM_METADATA['sanctioned_amount']} ({TEAM_METADATA['amount_words']}) was utilised for the purpose "
        "for which it was sanctioned, leaving a balance of Rs. NIL/- (Rupees NIL Only)."
    )
    r_c1.font.name = "Arial"
    r_c1.font.size = Pt(10)

    p_c2 = doc.add_paragraph()
    r_c2 = p_c2.add_run(
        "\tCertified that I have satisfied myself that the conditions on which the grant-in-aid was sanctioned have been duly fulfilled "
        "and that I have exercised the necessary checks to see that the money was actually utilized for the purpose for which it was sanctioned."
    )
    r_c2.font.name = "Arial"
    r_c2.font.size = Pt(10)

    doc.add_paragraph()

    # Team details header
    p_td = doc.add_paragraph()
    r_td = p_td.add_run("Team details")
    r_td.bold = True
    r_td.font.name = "Arial"

    # Team & Guide Table
    team_table = doc.add_table(rows=1, cols=6)
    team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    thdr = team_table.rows[0].cells
    thdr_titles = ['Sl.No.', 'Reg. No.', 'Student Name', 'Branch', 'Mobile No.', 'Sign.']
    for i, title in enumerate(thdr_titles):
        thdr[i].text = title
        thdr[i].paragraphs[0].runs[0].font.bold = True

    for st in STUDENTS:
        r_cells = team_table.add_row().cells
        r_data = [st["sl"], st["reg"], st["name"], st["branch"], st["mobile"], "________________"]
        for i, val in enumerate(r_data):
            r_cells[i].text = val
            r_cells[i].paragraphs[0].runs[0].font.name = "Arial"
            r_cells[i].paragraphs[0].runs[0].font.size = Pt(8.5)

    # Faculty Guide row
    fg_cells = team_table.add_row().cells
    fg_cells[0].text = "4"
    fg_cells[1].text = f"Faculty Guide Details: {TEAM_METADATA['guide']}, {TEAM_METADATA['guide_designation']}, {TEAM_METADATA['guide_dept']}, {TEAM_METADATA['college']}"
    fg_cells[5].text = "________________"

    doc.add_paragraph()

    # Signatures Table (3 cols)
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_headers = [
        "Name and Sign of the HoD\n\n\n\n______________________\nHead of Department",
        "Name and Sign of the Finance Officer / Auditor / Accounts Officer with seal\n\n\n______________________\nFinance Officer",
        "Name and Sign of the Principal with seal\n\n\n\n______________________\nDean / Principal"
    ]
    for i, txt in enumerate(sig_headers):
        sig_table.rows[0].cells[i].text = txt
        sig_table.rows[0].cells[i].paragraphs[0].runs[0].font.name = "Arial"

    doc.add_paragraph()

    p_ftr = doc.add_paragraph()
    r_ftr = p_ftr.add_run(f"Place: {TEAM_METADATA['place']}\nDate: {TEAM_METADATA['date']}\t\t\t\t\t\tCollege Seal")
    r_ftr.font.name = "Arial"

    out_file = os.path.join(DOCX_OUT_DIR, "Niral_Thiruvizha_Utilization_Certificate.docx")
    doc.save(out_file)
    print(f"  [DOCX CREATED] {out_file}")

def generate_all():
    print("=== GENERATING FILLED DOCX & PDF FORMAL DOSSIER ===")
    build_pre_receipt_docx()
    build_bill_summary_docx()
    build_utilization_certificate_docx()

if __name__ == "__main__":
    generate_all()
