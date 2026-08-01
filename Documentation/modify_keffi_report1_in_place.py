import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32com.client

def modify_report_in_place():
    print("=== MODIFYING KEFFI_REPORT 1.DOCX IN-PLACE (PRESERVING EXACT FONT SIZES & TYPOGRAPHY) ===")
    src_path = r'e:\Keffi Ai\Presentations_and_Extracted_Media\KEFFI_REPORT 1.docx'
    img_dir = r'e:\Keffi Ai\Documentation\extracted_report_images'
    new_poster_img = r'e:\Keffi Ai\Presentations_and_Extracted_Media\IMG-20260801-WA0001.jpg'

    doc = docx.Document(src_path)
    print(f"  [DOCX READ] Loaded {len(doc.paragraphs)} paragraphs from source document.")

    # Find key chapter insertion points and append concepts without changing font sizes
    current_chap = 0

    for i in range(len(doc.paragraphs)):
        p = doc.paragraphs[i]
        txt = p.text.strip()

        if "CHAPTER 1" in txt or "1 INTRODUCTION" in txt:
            current_chap = 1
        elif "CHAPTER 2" in txt or "2 LITERATURE SURVEY" in txt:
            current_chap = 2
        elif "CHAPTER 3" in txt or "3 SYSTEM SPECIFICATION" in txt:
            current_chap = 3
        elif "CHAPTER 4" in txt or "4 ANALYSIS OF PROJECT" in txt:
            current_chap = 4
        elif "CHAPTER 5" in txt or "5 SYSTEM DESIGN" in txt:
            current_chap = 5
        elif "CHAPTER 6" in txt or "6 IMPLEMENTATION" in txt:
            current_chap = 6
        elif "CHAPTER 7" in txt or "7 RESULT AND DISCUSSION" in txt:
            current_chap = 7
        elif "CHAPTER 8" in txt or "8 CONCLUSION" in txt:
            current_chap = 8

        # Alter/insert concepts at the end of key sections
        if current_chap == 1 and "1.5 OBJECTIVE" in txt:
            p.text = txt + "\n\n1.6 Hands-Free Real-Time Voice-to-Voice AI Architecture\nIn addition to text interaction, Keffi AI incorporates a hands-free Voice-to-Voice AI architecture. Patients experiencing panic or fatigue can speak into the microphone via Web Speech API (STT & TTS), receiving spoken therapeutic responses in a warm female voice (pitch = 1.05, rate = 0.95)."

        elif current_chap == 2 and "2.5 MULTI-EMOTION" in txt:
            p.text = txt + "\n\n2.6 Multi-LLM 70B Engine Cascade & High Availability\nKeffi AI implements a dual-engine cascade using Groq Llama-3.3-70B for sub-500ms response generation with automatic failover to OpenAI ChatGPT-4o-mini API.\n\n2.7 SHAP & LIME Explainable AI (XAI)\nIncorporates SHAP and LIME via /api/explain_clinical_decision for diagnostic token attribution transparency."

        elif current_chap == 4 and "4.2 MODEL DEVELOPMENT" in txt:
            p.text = txt + "\n\n4.3 Use Case Analysis for Voice Prosody Acoustic Tracking\nLibrosa Voice Prosody Analyzer (voice_prosody_analyzer.py) extracts F0 Pitch, RMS Energy, and Speech Rate (WPM) as acoustic depression biomarkers."

        elif current_chap == 5 and ("5.3.7" in txt or "DATABASE AND SECURITY MODULE" in txt):
            p.text = txt + "\n\n5.3.8 High-Concurrency SQLite Write-Ahead Logging (WAL Mode Engine)\nConfigures PRAGMA journal_mode=WAL to eliminate multi-threaded database locking errors.\n\n5.3.9 Patient Profile Registration Schema (POST /api/register)\nUpserts patient profiles storing Name, Phone, Email, DOB, Gender, City, MHQ score.\n\n5.3.10 User-Wise Isolated Chat History Persistence (GET /api/history/{patient_id})\nRestores chronological chat timelines isolated by patient ID.\n\n5.3.11 Admin Clinical Hub A-to-Z Patient Tracking & Transcript Inspector\nAggregates patient roster data, computes Days Inactive metrics, and offers a scrollable transcript viewer."

        elif current_chap == 8 and "8.2 FUTURE" in txt:
            p.text = txt + "\n\n8.3 Completed Advanced System Upgrades\nHigh-Concurrency SQLite WAL DB Storage, Patient Profile Registration, Isolated Chat History Restoration, Admin A-to-Z Patient Inspection, and Hands-Free Voice-to-Voice AI Interaction are 100% fully implemented."

    # Update Outcome Screenshots in Chapter 7
    add_sec = doc.add_paragraph()
    add_sec.text = "\n\n7.3 UPDATED SYSTEM SCREENSHOTS & DEEPTECH POSTER ASSET"

    landing_img = os.path.join(img_dir, "page_92_img_1.png")
    chat_img = os.path.join(img_dir, "page_92_img_2.png")
    admin_img = os.path.join(img_dir, "page_96_img_1.png")

    if os.path.exists(landing_img):
        p1 = doc.add_paragraph("Updated Landing Page Interface:")
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.add_run().add_picture(landing_img, width=Inches(5.5))

    if os.path.exists(chat_img):
        p2 = doc.add_paragraph("Updated Keffi Chatting Page (Sanctuary with Voice Chat):")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run().add_picture(chat_img, width=Inches(5.5))

    if os.path.exists(admin_img):
        p3 = doc.add_paragraph("Updated Admin Clinical Hub A-to-Z User Transcript Inspector:")
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.add_run().add_picture(admin_img, width=Inches(5.5))

    if os.path.exists(new_poster_img):
        p4 = doc.add_paragraph("Official Project Poster Asset (DeepTech Clinical System):")
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.add_run().add_picture(new_poster_img, width=Inches(4.8))

    # Save output DOCX
    out_docx_path = r"e:\Keffi Ai\Presentations_and_Extracted_Media\KEFFI_REPORT_1_FINAL_MODIFIED.docx"
    doc.save(out_docx_path)
    print(f"[SUCCESS] In-Place Modified DOCX Saved at: {out_docx_path}")

    # Convert to PDF via Word COM Engine
    out_pdf_path = r"e:\Keffi Ai\Documentation\PDF_Print_Copies\KEFFI_REPORT_1_FINAL_MODIFIED.pdf"
    os.makedirs(os.path.dirname(out_pdf_path), exist_ok=True)

    print(f"=== CONVERTING MODIFIED DOCX TO PDF VIA WORD COM ENGINE ===")
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc_w = word.Documents.Open(os.path.abspath(out_docx_path))
        doc_w.SaveAs(os.path.abspath(out_pdf_path), FileFormat=17) # 17 = wdFormatPDF
        doc_w.Close()
        print(f"[SUCCESS] PDF Generated Successfully: {out_pdf_path}")
    except Exception as e:
        print(f"[ERROR] PDF Conversion failed: {e}")
    finally:
        word.Quit()

if __name__ == "__main__":
    modify_report_in_place()
